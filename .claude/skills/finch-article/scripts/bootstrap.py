"""새 환경에서 스킬을 처음 사용할 때 실행하는 초기화 스크립트.

수행 작업:
  1. state/ 디렉토리 생성
  2. config.json 이 없으면 config.json.example 에서 복사
  3. ref/*.txt 가 없고 ref/*.docx 가 있으면 자동 추출 (python-docx 필요)
  4. 상태 보고

사용:
  python scripts/bootstrap.py
또는 자동 호출:
  python <pwd>/.claude/skills/finch-article/scripts/bootstrap.py
"""
import io, os, shutil, sys
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

SKILL_ROOT = Path(__file__).resolve().parent.parent
STATE = SKILL_ROOT / "state"
REF = SKILL_ROOT / "ref"
CONFIG = SKILL_ROOT / "config.json"
CONFIG_EXAMPLE = SKILL_ROOT / "config.json.example"

def ensure_state():
    STATE.mkdir(exist_ok=True)
    print(f"[ok] state dir: {STATE}")

def ensure_config():
    if CONFIG.exists():
        print(f"[ok] config: {CONFIG} (exists)")
    elif CONFIG_EXAMPLE.exists():
        shutil.copy(CONFIG_EXAMPLE, CONFIG)
        print(f"[new] config: copied from config.json.example")
        print(f"      → 필요하면 admin_url 등 수정하세요: {CONFIG}")
    else:
        print(f"[warn] config.json.example이 없어 건너뜀")

def ensure_ref_txt():
    if not REF.exists():
        print(f"[skip] ref dir not found")
        return
    txt_files = list(REF.glob("*.txt"))
    docx_files = list(REF.glob("*.docx"))
    if txt_files:
        print(f"[ok] ref txt files: {len(txt_files)}")
        return
    if not docx_files:
        print(f"[warn] no ref .txt or .docx files found")
        return
    try:
        from docx import Document  # python-docx
    except ImportError:
        print(f"[warn] python-docx 미설치 — 'pip install python-docx' 후 재실행하면 {len(docx_files)}개 docx를 txt로 추출합니다")
        return
    for dpath in docx_files:
        doc = Document(str(dpath))
        out = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t: out.append(t)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = p.text.strip()
                        if t: out.append(t)
        stem = dpath.stem.split("-")[0].strip()
        tpath = REF / f"{stem}.txt"
        tpath.write_text("\n".join(out), encoding="utf-8")
        print(f"[extract] {dpath.name} → {tpath.name} ({len(out)} blocks)")

if __name__ == "__main__":
    print(f"Skill root: {SKILL_ROOT}")
    ensure_state()
    ensure_config()
    ensure_ref_txt()
    print("\n다음 단계:")
    print("  1. Playwright 브라우저로 NS / SA / The Atlantic / finch admin 각각 수동 로그인")
    print("  2. /finch-article collect 로 수집 시작")
